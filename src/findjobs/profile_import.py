"""Local resume import and privacy-safe profile generation."""

import hashlib
import re
from pathlib import Path
from typing import Literal

from pydantic import BaseModel, ConfigDict, Field


# ---------------------------------------------------------------------------
# Pydantic models  (all frozen for immutability)
# ---------------------------------------------------------------------------


class ExperienceEntry(BaseModel):
    """A single work experience entry."""

    model_config = ConfigDict(frozen=True)

    title: str = ""
    company: str = ""
    start_date: str = ""
    end_date: str = ""
    description: str = ""


class ProjectEntry(BaseModel):
    """A single project entry."""

    model_config = ConfigDict(frozen=True)

    name: str = ""
    role: str = ""
    description: str = ""


class EducationEntry(BaseModel):
    """A single education entry."""

    model_config = ConfigDict(frozen=True)

    degree: str = ""
    institution: str = ""
    major: str = ""
    start_date: str = ""
    end_date: str = ""
    description: str = ""


class Profile(BaseModel):
    """Immutable privacy-safe resume profile.  No raw text or filename stored."""

    model_config = ConfigDict(frozen=True)

    schema_version: int = 1
    source_kind: Literal["docx", "pdf"]
    source_sha256: str
    contact_redacted: bool = True
    roles: list[str] = Field(default_factory=list)
    experience_years: float | None = None
    skills: list[str] = Field(default_factory=list)
    experiences: list[ExperienceEntry] = Field(default_factory=list)
    projects: list[ProjectEntry] = Field(default_factory=list)
    education: list[EducationEntry] = Field(default_factory=list)
    target_cities: list[str] = Field(default_factory=list)
    target_roles: list[str] = Field(default_factory=list)
    excluded_companies: list[str] = Field(default_factory=list)
    work_types: list[str] = Field(default_factory=list)
    constraints: list[str] = Field(default_factory=list)


# ---------------------------------------------------------------------------
# Canonical skill vocabulary — stable order, deduplicated
# ---------------------------------------------------------------------------

_CANONICAL_SKILLS: list[tuple[str, list[str]]] = [
    ("Python", [r"\bpython\b"]),
    ("Java", [r"\bjava\b"]),
    ("C", [r"\bC\b"]),
    ("C++", [r"c\+\+"]),
    ("Go", [r"\bGo\b"]),
    ("Rust", [r"\brust\b"]),
    ("AppSec", [r"appsec", r"application security", r"应用安全"]),
    ("security testing", [r"security test", r"安全测试", r"安全/合规测试"]),
    (
        "security test automation",
        [r"security test.*automation", r"安全测试.*自动化", r"自动化验证体系"],
    ),
    ("Harness engineering", [r"\bharness\b", r"测试harness"]),
    ("threat modeling", [r"threat model", r"威胁建模"]),
    ("code audit", [r"code audit", r"代码审计"]),
    ("IoT security", [r"iot\s*security", r"iot安全", r"智能家居.*安全"]),
    ("5G security", [r"5g.*安全", r"5g\s*security"]),
    ("penetration testing", [r"penetration test", r"pen test", r"渗透测试", r"渗透"]),
    (
        "vulnerability research",
        [r"vulnerability research", r"漏洞研究", r"漏洞挖掘", r"漏洞验证"],
    ),
    (
        "SDL",
        [r"\bsdl\b", r"security development lifecycle", r"安全开发生命周期", r"安全左移"],
    ),
    ("cloud security", [r"cloud security", r"云安全"]),
    ("data security", [r"data security", r"数据安全"]),
    (
        "privacy security",
        [r"privacy security", r"privacy compliance", r"隐私安全", r"隐私合规", r"隐私影响评估", r"\bprivacy\b"],
    ),
    ("GDPR", [r"\bgdpr\b"]),
    ("PIPL", [r"\bpipl\b", r"个人信息保护法"]),
    (
        "AI/LLM security",
        [r"ai security", r"llm security", r"ai\s*安全", r"llm\s*安全", r"ai.{0,8}安全", r"大模型安全"],
    ),
    ("Agent security", [r"agent security", r"agent\s*安全", r"智能体安全"]),
    ("prompt injection", [r"prompt injection", r"提示注入"]),
    (
        "tool-calling security",
        [r"tool.calling security", r"tool.use security", r"工具调用安全"],
    ),
    ("MLOps", [r"\bmlops\b"]),
    (
        "model deployment/fine-tuning",
        [r"model deployment", r"model fine.?tuning", r"模型部署", r"模型微调"],
    ),
]


def detect_skills(text: str) -> list[str]:
    """Return canonical skills matched in *text*, deduplicated, in stable order."""
    found: list[str] = []
    for canonical, patterns in _CANONICAL_SKILLS:
        for pat in patterns:
            if re.search(pat, text, re.IGNORECASE):
                found.append(canonical)
                break
    return found


# ---------------------------------------------------------------------------
# Canonical role vocabulary  (from explicit resume terms only)
# ---------------------------------------------------------------------------

_CANONICAL_ROLES: list[tuple[str, list[str]]] = [
    ("Security TSE", [r"security\s*tse", r"安全\s*tse"]),
    ("Security Engineer", [r"security engineer", r"安全工程师"]),
    ("Security Researcher", [r"security researcher", r"安全研究员", r"安全研究"]),
    ("Penetration Tester", [r"penetration tester", r"penetration test engineer", r"渗透测试"]),
    ("Security Architect", [r"security architect", r"安全架构师", r"安全架构"]),
    ("AppSec Engineer", [r"appsec engineer", r"application security engineer"]),
    ("SDL Engineer", [r"sdl engineer", r"sdl security"]),
    (
        "Vulnerability Researcher",
        [r"vulnerability researcher", r"漏洞研究员", r"漏洞研究"],
    ),
    ("Security Consultant", [r"security consultant", r"安全顾问"]),
    ("Software Engineer", [r"software engineer", r"软件工程师"]),
    ("DevOps Engineer", [r"devops engineer"]),
    ("MLOps Engineer", [r"mlops engineer"]),
    ("AI Engineer", [r"\bai engineer\b", r"ai engineer"]),
    ("Platform Engineer", [r"platform engineer", r"平台工程师"]),
    ("Infrastructure Engineer", [r"infrastructure engineer", r"基础架构工程师"]),
    ("Solution Architect", [r"solution architect", r"解决方案架构"]),
    ("Security Analyst", [r"security analyst", r"安全分析师", r"安全分析"]),
]


def _detect_roles(text: str) -> list[str]:
    """Return canonical roles matched in *text*, deduplicated, in stable order."""
    found: list[str] = []
    for canonical, patterns in _CANONICAL_ROLES:
        for pat in patterns:
            if re.search(pat, text, re.IGNORECASE):
                found.append(canonical)
                break
    return found


# ---------------------------------------------------------------------------
# Experience years detection  (explicit ``N+ years`` / ``N年以上`` only)
# ---------------------------------------------------------------------------

_EXPERIENCE_YEARS_RE = re.compile(r"(\d+)\+?\s*(?:years?|year|年|年以上)")


def _detect_experience_years(text: str) -> float | None:
    """Return the maximum explicit experience-year value found, or ``None``."""
    matches = [int(m) for m in _EXPERIENCE_YEARS_RE.findall(text)]
    return float(max(matches)) if matches else None


# ---------------------------------------------------------------------------
# Privacy redaction
# ---------------------------------------------------------------------------

_PHONE_RE = re.compile(r"1[3-9]\d{9}")
_EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
_ID_RE = re.compile(r"\b\d{17}[\dXx]\b")
_WECHAT_QQ_RE = re.compile(
    r"(?:微信|wechat|qq)\s*[:：]\s*\S+",
    re.IGNORECASE,
)


def _redact_pii(text: str, *, name: str = "") -> str:
    """Replace detected PII spans with ``[REDACTED]``."""
    text = _PHONE_RE.sub("[REDACTED]", text)
    text = _EMAIL_RE.sub("[REDACTED]", text)
    text = _ID_RE.sub("[REDACTED]", text)
    text = _WECHAT_QQ_RE.sub("[REDACTED]", text)
    if name:
        text = re.sub(re.escape(name), "[REDACTED]", text, flags=re.IGNORECASE)
    return text


# ---------------------------------------------------------------------------
# Conservative personal-name detection
#
# Return the first non-empty line that looks like a personal name:
#   - 2-6 CJK characters  (Chinese name)
#   - 2-4 title-cased Latin words  (Western name)
# Skip lines that look like section headings, contact labels, or are long.
# ---------------------------------------------------------------------------

# Lines that match these patterns are NOT names.
_SKIP_NAME_PATTERNS = [
    re.compile(p, re.IGNORECASE)
    for p in [
        # Section headings
        r"工作经验",
        r"工作经历",
        r"项目经验",
        r"项目经历",
        r"教育背景",
        r"教育经历",
        r"work experience",
        r"professional experience",
        r"employment",
        r"projects",
        r"education",
        r"background",
        r"summary",
        r"profile",
        r"skills?",
        r"技能",
        r"专业技能",
        r"自我评价",
        r"个人总结",
        r"关于我",
        # Contact labels
        r"手机|电话|phone|tel|mobile",
        r"邮箱|email|e-mail",
        r"微信|wechat",
        r"地址|address",
        r"linkedin",
        r"github",
    ]
]

# CJK name pattern: 2-6 CJK unified ideographs
_CJK_NAME = re.compile(r"^[一-鿿]{2,6}$")

# Western name pattern: 2-4 title-cased words (letters only, 2-20 chars each)
_WESTERN_NAME = re.compile(
    r"^[A-Z][a-z]+(?:[ '-][A-Z][a-z]+){1,3}$"
)


def _is_name_candidate(line: str) -> bool:
    """Return ``True`` if *line* looks like a plausible personal name."""
    stripped = line.strip()
    if not stripped:
        return False
    # Reject lines that match any skip pattern
    for pat in _SKIP_NAME_PATTERNS:
        if pat.search(stripped):
            return False
    # Reject long lines (names are short)
    if len(stripped) > 50:
        return False
    # CJK or Western name shape
    return bool(_CJK_NAME.match(stripped) or _WESTERN_NAME.match(stripped))


def _extract_likely_name(lines: list[str]) -> str:
    """Return the first line that looks like a personal name, or ``""``."""
    for line in lines:
        if _is_name_candidate(line.strip()):
            return line.strip()
    return ""


# ---------------------------------------------------------------------------
# DOCX extraction (python-docx)
# ---------------------------------------------------------------------------


def _extract_docx(path: Path) -> str:
    """Extract text via python-docx paragraphs and table cells."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX extraction.  Run: uv add python-docx"
        )

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ValueError(
            f"Failed to read DOCX file: {path}. The file may be corrupt."
        ) from exc

    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# PDF extraction (pypdf)
# ---------------------------------------------------------------------------


def _extract_pdf(path: Path) -> str:
    """Extract text via pypdf PdfReader (page-level text)."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError(
            "pypdf is required for PDF extraction.  Run: uv add pypdf"
        )

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise ValueError(
            f"Failed to read PDF file: {path}. The file may be corrupt."
        ) from exc

    if reader.is_encrypted:
        raise ValueError(f"PDF is encrypted or unreadable: {path}")

    parts: list[str] = []
    try:
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    except Exception as exc:
        raise ValueError(
            f"Failed to extract text from PDF file: {path}."
        ) from exc
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Section extraction  (conservative — no invented facts)
# ---------------------------------------------------------------------------

_SECTION_PATTERNS: dict[str, list[re.Pattern]] = {
    "experience": [
        re.compile(p, re.IGNORECASE)
        for p in [
            r"工作经验",
            r"工作经历",
            r"work experience",
            r"professional experience",
            r"employment",
        ]
    ],
    "projects": [
        re.compile(p, re.IGNORECASE)
        for p in [
            r"项目经验",
            r"项目经历",
            r"projects",
            r"project experience",
        ]
    ],
    "education": [
        re.compile(p, re.IGNORECASE)
        for p in [
            r"教育背景",
            r"教育经历",
            r"education",
            r"academic background",
        ]
    ],
}


def _classify_line(line: str) -> str | None:
    """Return section key (``experience``, ``projects``, ``education``)
    if *line* looks like a section heading, else ``None``."""
    for key, patterns in _SECTION_PATTERNS.items():
        for pat in patterns:
            if pat.search(line):
                return key
    return None


def _extract_sections(text: str) -> dict[str, list[str]]:
    """Split redacted text into sections by heading."""
    lines = [l.strip() for l in text.split("\n")]
    sections: dict[str, list[str]] = {
        "background": [],
        "experience": [],
        "projects": [],
        "education": [],
    }
    current = "background"
    for line in lines:
        if not line or line == "[REDACTED]":
            continue
        key = _classify_line(line)
        if key is not None:
            current = key
            continue
        sections[current].append(line)
    return sections


# ---------------------------------------------------------------------------
# SHA-256
# ---------------------------------------------------------------------------


def _file_sha256(path: Path) -> str:
    """Return the SHA-256 hex digest of *path* (filename is never stored)."""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        while True:
            chunk = f.read(65536)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()


# ---------------------------------------------------------------------------
# Markdown rendering
# ---------------------------------------------------------------------------


def render_markdown(profile: Profile) -> str:
    """Render *profile* as editable Markdown compatible with the project parser."""
    lines: list[str] = ["## Background", ""]
    lines.append("_Contact information has been redacted from this profile._")
    lines.append("")
    if profile.experience_years is not None:
        lines.append(f"- **Total experience**: {profile.experience_years} years")
    if profile.roles:
        lines.append(f"- **Roles**: {', '.join(profile.roles)}")
    if profile.skills:
        lines.append(f"- **Skills**: {', '.join(profile.skills)}")
    lines.append("")

    lines.append("## Target Cities")
    lines.append("")
    if profile.target_cities:
        for c in profile.target_cities:
            lines.append(f"- {c}")
    else:
        lines.append("_Not specified._")
    lines.append("")

    lines.append("## Target Roles")
    lines.append("")
    if profile.target_roles:
        for tr in profile.target_roles:
            lines.append(f"- {tr}")
    else:
        lines.append("_Not specified._")
    lines.append("")

    lines.append("## Salary Expectation")
    lines.append("")
    lines.append("_Not specified. Salary expectation must not be estimated._")
    lines.append("")

    lines.append("## Preferences")
    lines.append("")
    if profile.work_types:
        for wt in profile.work_types:
            lines.append(f"- {wt}")
    else:
        lines.append("_Not specified._")
    lines.append("")

    lines.append("## Excluded Companies")
    lines.append("")
    for ec in profile.excluded_companies:
        lines.append(f"- {ec}")
    lines.append("")

    lines.append("## Constraints")
    lines.append("")
    if profile.constraints:
        for c in profile.constraints:
            lines.append(f"- {c}")
    else:
        lines.append("_Not specified._")
    lines.append("")

    if profile.experiences:
        lines.append("## Work Experience")
        lines.append("")
        for e in profile.experiences:
            parts = []
            if e.title:
                parts.append(e.title)
            if e.start_date or e.end_date:
                parts.append(f"({e.start_date or '?'} – {e.end_date or '?'})")
            label = " ".join(parts)
            if e.company:
                lines.append(f"- **{e.company}**" + (f" — {label}" if label else ""))
            elif label:
                lines.append(f"- {label}")
            if e.description:
                lines.append(f"  - {e.description[:200]}")
        lines.append("")

    if profile.projects:
        lines.append("## Projects")
        lines.append("")
        for p in profile.projects:
            name = p.name or "Project"
            role = f" ({p.role})" if p.role else ""
            lines.append(f"- **{name}**{role}")
            if p.description:
                lines.append(f"  - {p.description[:200]}")
        lines.append("")

    if profile.education:
        lines.append("## Education")
        lines.append("")
        for e in profile.education:
            parts = []
            if e.degree:
                parts.append(e.degree)
            if e.major:
                parts.append(f"in {e.major}")
            if e.institution:
                parts.append(f"@ {e.institution}")
            if e.start_date or e.end_date:
                parts.append(f"({e.start_date or '?'} – {e.end_date or '?'})")
            label = " ".join(parts) if parts else (e.description[:200] if e.description else "")
            if label:
                lines.append(f"- {label}")
        lines.append("")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Rollback-safe pair writing
# ---------------------------------------------------------------------------


def _write_outputs_atomic(
    prepared: list[tuple[Path, str]],
    *,
    force: bool = False,
) -> None:
    """Write *prepared* (path, content) pairs atomically, with rollback.

    * All content is prepared before any write begins.
    * If ``force`` is set, existing files are backed up before overwrite.
    * On failure all original files are restored and temp/backup artifacts
      are cleaned up.  On success backups are removed.
    """
    backups: dict[Path, Path] = {}
    temp_files: list[tuple[Path, Path]] = []

    try:
        # -- back up existing outputs (when force) --------------------------
        for out_path, _content in prepared:
            if out_path.exists():
                if not force:
                    raise FileExistsError(
                        f"Output already exists: {out_path}.  Use --force to overwrite."
                    )
                backup = out_path.with_suffix(out_path.suffix + ".bak")
                out_path.rename(backup)
                backups[out_path] = backup

        # -- write temp siblings --------------------------------------------
        for out_path, content in prepared:
            out_path.parent.mkdir(parents=True, exist_ok=True)
            tmp = out_path.with_name(out_path.name + ".tmp")
            tmp.write_text(content, encoding="utf-8")
            temp_files.append((tmp, out_path))

        # -- replace --------------------------------------------------------
        for tmp, out_path in temp_files:
            tmp.replace(out_path)

    except BaseException:
        # Restore every backup regardless of destination state.
        for out_path, backup in backups.items():
            if backup.exists():
                if out_path.exists() and not out_path.is_dir():
                    out_path.unlink()
                if not out_path.is_dir():
                    backup.replace(out_path)
        # Remove any written-but-not-replaced temp files.
        for tmp, _out_path in temp_files:
            if tmp.exists():
                tmp.unlink()
        raise
    else:
        # clean up backups on success
        for backup in backups.values():
            if backup.exists():
                backup.unlink()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def import_resume(
    source: Path,
    *,
    force: bool = False,
    json_output: Path | None = None,
    markdown_output: Path | None = None,
) -> Profile:
    """Parse a DOCX/PDF resume into a privacy-safe ``Profile``.

    When *json_output* or *markdown_output* paths are provided the serialised
    profile is written atomically with rollback (see ``_write_outputs_atomic``).
    Raises ``FileExistsError`` for existing outputs unless *force* is set.
    """
    # ---- input validation (fast-fail before I/O) --------------------------
    if not source.exists():
        raise FileNotFoundError(f"Source file not found: {source}")
    if not source.is_file():
        raise ValueError(f"Source is not a regular file: {source}")
    ext = source.suffix.lower()
    if ext not in (".docx", ".pdf"):
        raise ValueError(f"Unsupported file extension: {ext!r} (expected .docx or .pdf)")

    # ---- collect what outputs the caller wants ----------------------------
    out_specs: list[tuple[str, Path]] = []
    for label, p in [("JSON", json_output), ("Markdown", markdown_output)]:
        if p is not None:
            out_specs.append((label, p))

    # ---- extract ----------------------------------------------------------
    text = extract_text(source)
    sha256 = _file_sha256(source)

    # ---- redact -----------------------------------------------------------
    lines = text.splitlines()
    name = _extract_likely_name(lines)
    redacted = _redact_pii(text, name=name)
    # Replace the first occurrence of the name line with placeholder
    redacted_lines = redacted.splitlines()
    if redacted_lines and name:
        for i, rl in enumerate(redacted_lines):
            if rl.strip() and name.lower() in rl.lower():
                redacted_lines[i] = "[REDACTED]"
                break
    redacted = "\n".join(redacted_lines)

    # ---- structure --------------------------------------------------------
    sections = _extract_sections(redacted)
    skills = detect_skills(redacted)
    roles = _detect_roles(redacted)
    experience_years = _detect_experience_years(redacted)

    experiences = [ExperienceEntry(description=l[:200]) for l in sections["experience"] if l]
    projects = [ProjectEntry(description=l[:200]) for l in sections["projects"] if l]
    education = [EducationEntry(description=l[:200]) for l in sections["education"] if l]

    profile = Profile(
        source_kind="docx" if ext == ".docx" else "pdf",
        source_sha256=sha256,
        contact_redacted=True,
        skills=skills,
        roles=roles,
        experience_years=experience_years,
        experiences=experiences,
        projects=projects,
        education=education,
        excluded_companies=["huawei"],
    )

    # ---- atomic write (rollback-safe) ------------------------------------
    prepared: list[tuple[Path, str]] = []
    for label, out_path in out_specs:
        if label == "JSON":
            content = profile.model_dump_json(indent=2) + "\n"
        else:
            content = render_markdown(profile)
        prepared.append((out_path, content))

    if prepared:
        _write_outputs_atomic(prepared, force=force)

    return profile


def extract_text(path: Path) -> str:
    """Extract raw text from *path* (``.docx`` or ``.pdf``)."""
    ext = path.suffix.lower()
    if ext == ".docx":
        text = _extract_docx(path)
    elif ext == ".pdf":
        text = _extract_pdf(path)
    else:
        raise ValueError(f"Unsupported file extension: {ext!r} (expected .docx or .pdf)")

    stripped = text.strip()
    if not stripped:
        raise ValueError(f"Extracted text is empty: {path}")
    return stripped
