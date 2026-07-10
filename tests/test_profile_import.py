"""Tests for local resume import (Phase 3A)."""

import json
import re
from pathlib import Path
from unittest.mock import patch

import pytest
from pydantic import ValidationError

from findjobs.profile_import import (
    _classify_line,
    _detect_experience_years,
    _detect_roles,
    _extract_likely_name,
    _extract_sections,
    _file_sha256,
    _is_name_candidate,
    _redact_pii,
    detect_skills,
    EducationEntry,
    ExperienceEntry,
    import_resume,
    Profile,
    ProjectEntry,
    render_markdown,
)

# ---------------------------------------------------------------------------
# Helpers / internals
# ---------------------------------------------------------------------------


class TestIsNameCandidate:
    def test_cjk_name(self):
        assert _is_name_candidate("张三") is True
        assert _is_name_candidate("欧阳春雪") is True
        assert _is_name_candidate("张") is False  # too short
        assert _is_name_candidate("一二三四五六七八") is False  # too long (8 CJK)

    def test_western_name(self):
        assert _is_name_candidate("John Doe") is True
        assert _is_name_candidate("Alice Bob Charlie") is True
        assert _is_name_candidate("J") is False  # single word
        assert _is_name_candidate("A B C D E") is False  # 5 words

    def test_heading_rejected(self):
        assert _is_name_candidate("Work Experience") is False
        assert _is_name_candidate("Projects") is False
        assert _is_name_candidate("Education") is False
        assert _is_name_candidate("技能") is False
        assert _is_name_candidate("项目经验") is False

    def test_contact_label_rejected(self):
        assert _is_name_candidate("Phone: 13812345678") is False
        assert _is_name_candidate("Email: test@foo.com") is False
        assert _is_name_candidate("微信: mywechat") is False

    def test_long_line_rejected(self):
        assert _is_name_candidate("A" * 60) is False


class TestExtractLikelyName:
    def test_cjk_name_detected(self):
        assert _extract_likely_name(["张三", "other"]) == "张三"

    def test_western_name_detected(self):
        assert _extract_likely_name(["", "John Doe", "Engineer"]) == "John Doe"

    def test_heading_at_start_is_not_name(self):
        """Section headings at the first line are not mistaken for names."""
        assert _extract_likely_name(["## Projects", "Content"]) == ""
        assert _extract_likely_name(["Work Experience", "Engineer"]) == ""
        assert _extract_likely_name(["Contact: 13812345678"]) == ""

    def test_empty_lines_returns_empty(self):
        assert _extract_likely_name(["", "  "]) == ""

    def test_no_name_returns_empty(self):
        assert _extract_likely_name(["## Projects", "some content"]) == ""


class TestClassifyLine:
    def test_chinese_headings(self):
        assert _classify_line("工作经历") == "experience"
        assert _classify_line("项目经验") == "projects"
        assert _classify_line("教育背景") == "education"

    def test_english_headings(self):
        assert _classify_line("Work Experience") == "experience"
        assert _classify_line("PROJECTS") == "projects"
        assert _classify_line("Education") == "education"


class TestExtractSections:
    def test_lines_go_to_correct_section(self):
        text = (
            "张三\n"
            "test@example.com\n"
            "## Work Experience\n"
            "Senior Engineer at FooCorp\n"
            "## Projects\n"
            "Project Alpha\n"
            "## Education\n"
            "Master of CS\n"
        )
        sections = _extract_sections(text)
        assert any("Senior Engineer" in l for l in sections["experience"])
        assert any("Project Alpha" in l for l in sections["projects"])
        assert any("Master of CS" in l for l in sections["education"])


class TestRedactPii:
    _SAMPLE = (
        "My name is 张三.\n"
        "Phone: 13812345678\n"
        "Email: test@example.com\n"
        "WeChat: mywechat\n"
        "ID: 110101199001011234\n"
    )

    def test_redacts_phone(self):
        result = _redact_pii(self._SAMPLE)
        assert "13812345678" not in result
        assert "[REDACTED]" in result

    def test_redacts_email(self):
        result = _redact_pii(self._SAMPLE)
        assert "test@example.com" not in result

    def test_redacts_wechat_qq(self):
        result = _redact_pii(self._SAMPLE)
        assert "mywechat" not in result

    def test_redacts_national_id(self):
        result = _redact_pii(self._SAMPLE)
        assert "110101199001011234" not in result

    def test_redacts_name_when_provided(self):
        result = _redact_pii("张三 is my name.", name="张三")
        assert "张三" not in result

    def test_no_name_no_redact(self):
        result = _redact_pii("Keep this name", name="")
        assert "Keep this name" in result


class TestFileSha256:
    def test_sha256_no_filename(self, tmp_path):
        f = tmp_path / "resume.docx"
        f.write_bytes(b"hello")
        digest = _file_sha256(f)
        assert re.fullmatch(r"[0-9a-f]{64}", digest)
        # digest should be content-based, not path-based
        assert "resume.docx" not in digest


class TestDetectSkills:
    def test_detects_python(self):
        assert "Python" in detect_skills("I write Python code")

    def test_detects_rust(self):
        assert "Rust" in detect_skills("Rust systems programming")

    def test_dedup_and_stable_order(self):
        skills = detect_skills(
            "Python and Rust and Python again and Java and Rust"
        )
        # Ordered by CANONICAL_SKILLS: Python, Java, ... Rust
        assert skills == ["Python", "Java", "Rust"]

    def test_empty_text_returns_empty(self):
        assert detect_skills("") == []

    def test_detects_resume_security_engineering_phrases(self):
        skills = detect_skills(
            "AI驱动的安全测试Harness开发、漏洞验证、隐私影响评估、"
            "安全左移、IoT安全与5G行业安全"
        )

        assert {
            "security testing",
            "Harness engineering",
            "vulnerability research",
            "privacy security",
            "SDL",
            "IoT security",
            "5G security",
            "AI/LLM security",
        }.issubset(skills)


class TestDetectRoles:
    def test_detects_security_engineer(self):
        assert "Security Engineer" in _detect_roles("Security Engineer at FooCorp")

    def test_dedup_and_order(self):
        roles = _detect_roles("Penetration Tester and Security Engineer")
        assert roles == ["Security Engineer", "Penetration Tester"]

    def test_empty_returns_empty(self):
        assert _detect_roles("") == []


class TestDetectExperienceYears:
    def test_detects_years(self):
        assert _detect_experience_years("5+ years experience") == 5.0

    def test_detects_chinese_years(self):
        assert _detect_experience_years("8年以上工作经验") == 8.0

    def test_takes_maximum(self):
        assert _detect_experience_years("3 years then 10+ years") == 10.0

    def test_no_match_returns_none(self):
        assert _detect_experience_years("no years mentioned") is None


# ====================  Pydantic models (frozen)  ===========================


class TestProfileModel:
    def test_defaults(self):
        p = Profile(source_kind="docx", source_sha256="abcd" * 16)
        assert p.schema_version == 1
        assert p.contact_redacted is True
        assert p.excluded_companies == []
        assert p.skills == []

    def test_frozen_assignment_raises(self):
        p = Profile(source_kind="pdf", source_sha256="abcd" * 16)
        with pytest.raises(ValidationError):
            p.skills = ["Python"]

    def test_excluded_companies_set(self):
        p = Profile(
            source_kind="docx",
            source_sha256="abcd" * 16,
            excluded_companies=["huawei"],
        )
        assert p.excluded_companies == ["huawei"]


class TestExperienceEntryFrozen:
    def test_frozen(self):
        e = ExperienceEntry()
        with pytest.raises(ValidationError):
            e.title = "changed"


class TestProjectEntryFrozen:
    def test_frozen(self):
        p = ProjectEntry()
        with pytest.raises(ValidationError):
            p.name = "changed"


class TestEducationEntryFrozen:
    def test_frozen(self):
        e = EducationEntry()
        with pytest.raises(ValidationError):
            e.degree = "changed"


# ====================  DOCX extraction (realistic fake)  ===================


class TestDocxExtraction:
    def test_paragraphs_and_tables(self, tmp_path):
        """DOCX paragraphs and table cells are both extracted (real document)."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("张三")
        doc.add_paragraph("Work Experience")
        doc.add_paragraph("Engineer")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Skill: Python"
        table.cell(0, 1).text = "Skill: Rust"

        docx_path = tmp_path / "resume.docx"
        doc.save(str(docx_path))

        from findjobs.profile_import import _extract_docx

        text = _extract_docx(docx_path)
        assert "Engineer" in text
        assert "Skill: Python" in text

    def test_corrupt_docx_raises(self, tmp_path):
        """Corrupt/unreadable docx raises ValueError with descriptive message."""
        docx_path = tmp_path / "corrupt.docx"
        docx_path.write_bytes(b"\x00\x01\x02")

        def _raise(_path):
            raise Exception("corrupt file")

        with patch("docx.Document", side_effect=_raise):
            from findjobs.profile_import import _extract_docx

            with pytest.raises(ValueError, match="corrupt"):
                _extract_docx(docx_path)

    def test_missing_docx_file(self, tmp_path):
        """Missing source raises FileNotFoundError."""
        missing = tmp_path / "nonexistent.docx"
        with pytest.raises(FileNotFoundError):
            import_resume(missing)


class TestPdfExtraction:
    def test_encrypted_pdf_raises(self, tmp_path):
        """Encrypted PDF raises ValueError."""
        pdf_path = tmp_path / "secret.pdf"
        pdf_path.write_bytes(b"%PDF-encrypted")

        with patch("findjobs.profile_import._extract_pdf") as mock_extract:
            mock_extract.side_effect = ValueError("PDF is encrypted or unreadable")
            with pytest.raises(ValueError, match="encrypted"):
                import_resume(pdf_path)

    def test_successful_pdf(self, tmp_path):
        """Successful PDF extraction returns text."""
        pdf_path = tmp_path / "resume.pdf"
        pdf_path.write_bytes(b"%PDF-1.4")

        def fake_extract(path):
            return "John Doe\nWork Experience\nDeveloper at Corp"

        with patch("findjobs.profile_import._extract_pdf", side_effect=fake_extract):
            profile = import_resume(pdf_path)
            assert profile.source_kind == "pdf"
            assert profile.source_sha256

    def test_blank_pdf_raises(self, tmp_path):
        """Blank/unreadable PDF (empty extracted text) raises ValueError."""
        pdf_path = tmp_path / "blank.pdf"
        pdf_path.write_bytes(b"%PDF-1.4")

        with patch(
            "findjobs.profile_import._extract_pdf",
            return_value="   ",
        ):
            with pytest.raises(ValueError, match="empty"):
                import_resume(pdf_path)


# ====================  Unsupported / missing input  ========================


class TestUnsupportedInput:
    def test_unsupported_extension(self, tmp_path):
        bad = tmp_path / "resume.txt"
        bad.write_text("hello")
        with pytest.raises(ValueError, match="Unsupported file extension"):
            import_resume(bad)

    def test_missing_file(self, tmp_path):
        missing = tmp_path / "missing.docx"
        with pytest.raises(FileNotFoundError):
            import_resume(missing)


# ====================  Section extraction  =================================


class TestSectionExtraction:
    def test_experience_extracted(self, tmp_path):
        """Lines under a Work Experience heading become experience entries."""
        docx_path = tmp_path / "resume.docx"
        docx_path.write_bytes(b"fake")

        with patch(
            "findjobs.profile_import._extract_docx",
            return_value="Name\n## Work Experience\nEngineer at Foo\nLed team\n## Projects\nSomething",
        ):
            profile = import_resume(docx_path)
            assert len(profile.experiences) >= 1
            assert profile.experiences[0].description

    def test_projects_extracted(self, tmp_path):
        """Projects section is preserved (heading not mistaken for name)."""
        docx_path = tmp_path / "resume.docx"
        docx_path.write_bytes(b"fake")

        with patch(
            "findjobs.profile_import._extract_docx",
            return_value="## Projects\nBuilt X\nBuilt Y",
        ):
            profile = import_resume(docx_path)
            assert len(profile.projects) >= 1

    def test_education_extracted(self, tmp_path):
        """Education section is preserved (heading not mistaken for name)."""
        docx_path = tmp_path / "resume.docx"
        docx_path.write_bytes(b"fake")

        with patch(
            "findjobs.profile_import._extract_docx",
            return_value="## Education\nMaster of CS\nBachelor of Engineering",
        ):
            profile = import_resume(docx_path)
            assert len(profile.education) >= 1

    def test_sections_preserve_experience_years(self, tmp_path):
        """Explicit experience-years are extracted and stored."""
        docx_path = tmp_path / "resume.docx"
        docx_path.write_bytes(b"fake")

        with patch(
            "findjobs.profile_import._extract_docx",
            return_value="John Doe\n## Work Experience\nEngineer (5+ years)",
        ):
            profile = import_resume(docx_path)
            assert profile.experience_years == 5.0

    def test_roles_extracted(self, tmp_path):
        """Canonical roles are extracted from resume text."""
        docx_path = tmp_path / "resume.docx"
        docx_path.write_bytes(b"fake")

        with patch(
            "findjobs.profile_import._extract_docx",
            return_value="John Doe\nSecurity Engineer\nPenetration Tester",
        ):
            profile = import_resume(docx_path)
            assert "Security Engineer" in profile.roles


# ====================  Privacy / redaction  ================================


class TestPrivacyRedaction:
    @staticmethod
    def _run_import(docx_text: str, tmp_path: Path) -> Profile:
        path = tmp_path / "resume.docx"
        path.write_bytes(b"fake")
        with patch(
            "findjobs.profile_import._extract_docx",
            return_value=docx_text,
        ):
            return import_resume(path, json_output=tmp_path / "out.json")

    def test_name_redacted_in_json(self, tmp_path):
        """Likely first-line name is redacted from JSON output."""
        self._run_import("张三\nWork Experience\nEngineer", tmp_path)
        j = json.loads((tmp_path / "out.json").read_text(encoding="utf-8"))
        assert "张三" not in json.dumps(j)

    def test_phone_redacted(self, tmp_path):
        self._run_import("John\nPhone: 13812345678\n", tmp_path)
        j = json.loads((tmp_path / "out.json").read_text(encoding="utf-8"))
        assert "13812345678" not in json.dumps(j)

    def test_email_redacted(self, tmp_path):
        self._run_import("John\nEmail: test@foo.com\n", tmp_path)
        j = json.loads((tmp_path / "out.json").read_text(encoding="utf-8"))
        assert "test@foo.com" not in json.dumps(j)

    def test_markdown_has_no_pii(self, tmp_path):
        """Markdown output does not contain the extracted PII."""
        path = tmp_path / "resume.docx"
        path.write_bytes(b"fake")
        with patch(
            "findjobs.profile_import._extract_docx",
            return_value="张三\nEmail: test@foo.com\n",
        ):
            import_resume(path, markdown_output=tmp_path / "out.md")
        md = (tmp_path / "out.md").read_text(encoding="utf-8")
        assert "张三" not in md
        assert "test@foo.com" not in md

    def test_sha_has_no_filename(self, tmp_path):
        """SHA256 does not contain any filename reference."""
        path = tmp_path / "my_resume.docx"
        path.write_bytes(b"some content")
        sha = _file_sha256(path)
        assert "my_resume" not in sha
        assert "docx" not in sha


# ====================  Overwrite refusal / force  ==========================


class TestOverwriteProtection:
    def test_refuses_overwrite_json(self, tmp_path):
        src = tmp_path / "resume.docx"
        src.write_bytes(b"fake")
        out = tmp_path / "profile.json"
        out.write_text("{}")

        with patch(
            "findjobs.profile_import._extract_docx",
            return_value="Name\nSkills",
        ):
            with pytest.raises(FileExistsError, match="already exists"):
                import_resume(src, json_output=out)

    def test_force_overwrites(self, tmp_path):
        src = tmp_path / "resume.docx"
        src.write_bytes(b"fake")
        out = tmp_path / "profile.json"
        out.write_text("old")

        with patch(
            "findjobs.profile_import._extract_docx",
            return_value="John Doe\nSkills",
        ):
            profile = import_resume(src, json_output=out, force=True)
        assert json.loads(out.read_text(encoding="utf-8"))["skills"] == []


# ====================  Atomic / rollback-safe write  =======================


class TestAtomicWrite:
    def test_render_failure_writes_nothing(self, tmp_path):
        """A render failure before write stage leaves no output files."""
        src = tmp_path / "resume.docx"
        src.write_bytes(b"fake")
        json_out = tmp_path / "profile.json"
        md_out = tmp_path / "profile.md"

        with patch(
            "findjobs.profile_import._extract_docx",
            return_value="Name\nHello",
        ):
            with patch(
                "findjobs.profile_import.render_markdown",
                side_effect=RuntimeError("md fail"),
            ):
                with pytest.raises(RuntimeError):
                    import_resume(src, json_output=json_out, markdown_output=md_out)

        assert not json_out.exists()
        assert not md_out.exists()
        assert list(tmp_path.glob("*.tmp")) == []

    def test_second_write_failure_restores_originals(self, tmp_path):
        """If the second replace fails, both originals are restored."""
        import pathlib

        src = tmp_path / "resume.docx"
        src.write_bytes(b"fake")
        json_out = tmp_path / "profile.json"
        md_out = tmp_path / "profile.md"

        json_out.write_text("original json content")
        md_out.write_text("original md content")

        _real_replace = pathlib.Path.replace

        def _fail_second_replace(self, target):
            # Fail only when a .tmp file is being replaced onto md_out.
            if self.name.endswith(".tmp") and target == md_out:
                raise OSError("simulated md replace failure")
            return _real_replace(self, target)

        with patch(
            "findjobs.profile_import._extract_docx",
            return_value="John Doe\nHello",
        ):
            with patch.object(Path, "replace", _fail_second_replace):
                with pytest.raises(OSError):
                    import_resume(
                        src,
                        json_output=json_out,
                        markdown_output=md_out,
                        force=True,
                    )

        assert json_out.read_text() == "original json content"
        assert md_out.read_text() == "original md content"
        assert list(tmp_path.glob("*.tmp")) == []
        assert list(tmp_path.glob("*.bak")) == []


# ====================  CLI integration  ====================================


class TestCliIntegration:
    def test_cli_success_with_defaults(self, tmp_path):
        """CLI import with omittable defaults succeeds."""
        from typer.testing import CliRunner
        from findjobs.cli import app

        src = tmp_path / "resume.docx"
        src.write_bytes(b"fake")

        with patch(
            "findjobs.profile_import._extract_docx",
            return_value="John Doe\nWork Experience\nEngineer",
        ):
            runner = CliRunner()
            # Use explicit paths to avoid writing to real profile/
            result = runner.invoke(
                app,
                [
                    "profile",
                    "import",
                    str(src),
                    "--json-output",
                    str(tmp_path / "profile.json"),
                    "--markdown-output",
                    str(tmp_path / "profile.md"),
                ],
            )
        assert result.exit_code == 0, result.output
        assert (tmp_path / "profile.json").exists()
        assert (tmp_path / "profile.md").exists()

    def test_cli_error_exits_nonzero(self, tmp_path):
        """Missing file produces nonzero CLI exit."""
        from typer.testing import CliRunner
        from findjobs.cli import app

        runner = CliRunner()
        result = runner.invoke(
            app,
            [
                "profile",
                "import",
                str(tmp_path / "nonexistent.docx"),
                "--json-output",
                str(tmp_path / "out.json"),
                "--markdown-output",
                str(tmp_path / "out.md"),
            ],
        )
        assert result.exit_code != 0


def test_json_keeps_required_nullable_profile_fields(tmp_path):
    source = tmp_path / "resume.docx"
    source.write_bytes(b"fake")
    output = tmp_path / "profile.json"

    with patch(
        "findjobs.profile_import._extract_docx",
        return_value="Security Engineer\n安全测试",
    ):
        import_resume(source, json_output=output)

    payload = json.loads(output.read_text(encoding="utf-8"))
    assert "experience_years" in payload
    assert payload["experience_years"] is None


# ====================  No database access  =================================


class TestNoDatabaseAccess:
    def test_no_db_in_profile_import_module(self):
        """The profile_import module does not import findjobs.db."""
        import findjobs.profile_import

        source = Path(findjobs.profile_import.__file__).read_text(encoding="utf-8")
        assert "findjobs.db" not in source
        assert "from findjobs.db" not in source
